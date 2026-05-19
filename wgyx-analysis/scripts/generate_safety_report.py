#!/usr/bin/env python3
"""Generate a Chinese safety-production violation analysis Markdown report from an Excel export."""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd

# 用于生成 Word 报告
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 用于生成 PDF 报告
from fpdf import FPDF


STATUS_COL = "状态"
AREA_COL = "违章区域"
INSPECT_ORG_COL = "检查机构"
INSPECTOR_TYPE_COL = "检查人员类别"
TIME_COL = "违章时间"
PERSON_TYPE_COL = "违章人员性质"
GENERAL_CONTRACTOR_COL = "总包单位"
SUBCONTRACTOR_COL = "分包单位"
VIOLATION_DESC_COL = "违章情况"
VIOLATION_TYPE_COL = "违章类别"
VIOLATION_LEVEL_COL = "违章级别"
SCORE_COL = "分值"
SOURCE_COL = "数据来源"
CLAUSE_COL = "条款内容"


def setup_chinese_font() -> None:
    candidates = [
        "Microsoft YaHei",
        "SimHei",
        "Noto Sans CJK SC",
        "Noto Sans CJK JP",
        "WenQuanYi Micro Hei",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["font.sans-serif"] = candidates
    plt.rcParams["axes.unicode_minus"] = False


def read_export(path: Path) -> tuple[pd.DataFrame, str]:
    raw = pd.read_excel(path, sheet_name=0, header=None)
    period = ""
    for value in raw.stack().dropna().astype(str).head(50):
        if "时间范围" in value:
            period = value.replace("时间范围：", "").replace("时间范围:", "").strip()
            break

    header_idx = None
    for i in range(min(20, len(raw))):
        row = [str(x).strip() for x in raw.iloc[i].fillna("").tolist()]
        if STATUS_COL in row and VIOLATION_DESC_COL in row:
            header_idx = i
            break
    if header_idx is None:
        raise ValueError("无法识别表头行，请确认 Excel 包含“状态”“违章情况”等字段。")

    df = pd.read_excel(path, sheet_name=0, header=header_idx)
    df = df.loc[:, ~df.columns.astype(str).str.startswith("Unnamed")]
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all")
    if STATUS_COL in df.columns:
        df = df[df[STATUS_COL].astype(str).str.strip().ne("已删除")]
    if TIME_COL in df.columns:
        df[TIME_COL] = pd.to_datetime(df[TIME_COL], errors="coerce")
    if SCORE_COL in df.columns:
        df[SCORE_COL] = pd.to_numeric(df[SCORE_COL], errors="coerce").fillna(0)
    return df, period


def norm_series(df: pd.DataFrame, col: str) -> pd.Series:
    if col not in df.columns:
        return pd.Series(dtype=object)
    return (
        df[col]
        .fillna("未填写")
        .astype(str)
        .str.strip()
        .replace({"": "未填写", "nan": "未填写"})
    )


def shorten_area(text: str) -> str:
    parts = [p for p in str(text).split("/") if p]
    if len(parts) >= 3:
        return "/".join(parts[-3:])
    return str(text)


def classify_topic(row: pd.Series) -> str:
    text = " ".join(str(row.get(c, "")) for c in [VIOLATION_DESC_COL, CLAUSE_COL, VIOLATION_TYPE_COL])
    rules = [
        ("挂牌/能量隔离", r"挂牌|摘牌|锁定|联锁|能量|确认"),
        ("个体防护/PPE", r"防护|护目镜|面罩|耳塞|口罩|手套|安全带|围裙|劳防"),
        ("高处/临边", r"高处|登高|临边|孔洞|坠落|安全绳|护栏"),
        ("动火/消防", r"动火|火星|灭火器|接火|可燃|氧气|乙炔|消防"),
        ("起重吊装", r"起重|吊装|吊物|吊索|司索|指吊|行车|止摆|斜吊"),
        ("安全交底/工票", r"交底|工票|签字|代签|审批|许可"),
        ("设备设施/机械伤害", r"设备|叉车|小车|步进梁|打捆|机械|旋转|挤压"),
        ("行为规范/禁令", r"吸烟|酒后|手机|禁令|无证|准入|睡岗"),
    ]
    for label, pattern in rules:
        if re.search(pattern, text):
            return label
    return str(row.get(VIOLATION_TYPE_COL, "其他")) or "其他"


def save_bar(series: pd.Series, title: str, path: Path, xlabel: str = "数量") -> None:
    setup_chinese_font()
    data = series.dropna().head(10).sort_values()
    fig, ax = plt.subplots(figsize=(8, max(4, 0.45 * len(data))))
    ax.barh(data.index.astype(str), data.values)
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    for i, v in enumerate(data.values):
        ax.text(v, i, f" {int(v)}", va="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def save_pie(series: pd.Series, title: str, path: Path) -> None:
    setup_chinese_font()
    data = series.dropna()
    fig, ax = plt.subplots(figsize=(6.2, 4.6))
    ax.pie(data.values, labels=data.index.astype(str), autopct="%1.1f%%", startangle=90)
    ax.set_title(title)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def save_trend(df: pd.DataFrame, path: Path) -> pd.Series:
    setup_chinese_font()
    if TIME_COL not in df.columns or df[TIME_COL].isna().all():
        trend = pd.Series(dtype=int)
    else:
        trend = df.dropna(subset=[TIME_COL]).set_index(TIME_COL).resample("W").size()

    fig, ax = plt.subplots(figsize=(8, 4.2))
    if len(trend):
        ax.plot(trend.index, trend.values, marker="o")
        ax.set_xticks(trend.index)
        ax.set_xticklabels([d.strftime("%m-%d") for d in trend.index], rotation=35, ha="right")
    ax.set_title("违章发生趋势（按周）")
    ax.set_ylabel("违章数量")
    ax.grid(True, alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return trend


def format_count_table(series: pd.Series, value_name: str = "数量") -> list[str]:
    lines = [f"| 项目 | {value_name} | 占比 |", "| --- | ---: | ---: |"]
    total = float(series.sum()) or 1.0
    for idx, val in series.head(10).items():
        numeric = float(val)
        value = str(int(numeric)) if numeric.is_integer() else f"{numeric:.1f}"
        lines.append(f"| {idx} | {value} | {numeric / total * 100:.1f}% |")
    return lines


def relative_asset_path(output_path: Path, asset_path: Path) -> str:
    return asset_path.relative_to(output_path.parent).as_posix()


def export_stats(excel_path: Path) -> None:
    df, period = read_export(excel_path)
    if df.empty:
        print(json.dumps({"error": "清洗后没有有效违章记录"}, ensure_ascii=False))
        return

    df["专题归类"] = df.apply(classify_topic, axis=1)
    df["简化区域"] = norm_series(df, AREA_COL).map(shorten_area)

    level_counts = norm_series(df, VIOLATION_LEVEL_COL).value_counts()
    topic_counts = norm_series(df, "专题归类").value_counts()
    area_counts = norm_series(df, "简化区域").value_counts()

    total = len(df)
    score_total = df[SCORE_COL].sum() if SCORE_COL in df.columns else 0
    a_count = int(level_counts.get("A类违章", 0) or level_counts.get("A类", 0) or 0)

    stats = {
        "period": period,
        "total_records": total,
        "total_score": float(score_total),
        "a_level_count": a_count,
        "top_topics": {str(k): int(v) for k, v in topic_counts.head(5).items()},
        "top_areas": {str(k): int(v) for k, v in area_counts.head(5).items()},
    }
    print(json.dumps(stats, ensure_ascii=False, indent=2))


def set_run_font(run, font_name="微软雅黑", size_pt=10.5, color_rgb=None, bold=False):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    run.bold = bold


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_docx_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, font_name="微软雅黑", size_pt=10, bold=True, color_rgb=RGBColor(27, 54, 93))
                
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            set_cell_margins(row_cells[i], top=100, bottom=100)
            for paragraph in row_cells[i].paragraphs:
                if i == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, font_name="微软雅黑", size_pt=9.5)
                    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build_docx(docx_path: Path, title: str, period: str, total: int, score_total: float, a_count: int, high_ratio: float, 
               level_counts: pd.Series, topic_counts: pd.Series, area_counts: pd.Series, person_counts: pd.Series, source_counts: pd.Series,
               root_causes: list[str], actions: list[dict], 
               trend_path: Path, level_path: Path, topic_path: Path, area_path: Path, person_path: Path, source_path: Path) -> None:
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(title)
    set_run_font(run_title, font_name="微软雅黑", size_pt=20, bold=True, color_rgb=RGBColor(27, 54, 93))
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(24)
    meta_text = (
        f"报告日期：{pd.Timestamp.today().strftime('%Y年%m月%d日')}   |   "
        f"数据周期：{period}\n"
        f"数据口径：共 {total} 条有效违章记录，剔除“已删除”记录；"
        f"累计扣分约 {score_total:.0f} 分，A类/重大违章 {a_count} 起（占比 {high_ratio:.1f}%）。"
    )
    run_meta = p_meta.add_run(meta_text)
    set_run_font(run_meta, font_name="微软雅黑", size_pt=10, color_rgb=RGBColor(120, 120, 120))
    
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, font_name="微软雅黑", size_pt=14, bold=True, color_rgb=RGBColor(27, 54, 93))
        
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, font_name="微软雅黑", size_pt=12, bold=True, color_rgb=RGBColor(27, 54, 93))

    def add_paragraph(text, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Inches(0.3)
        run = p.add_run(text)
        set_run_font(run, font_name="微软雅黑", size_pt=10.5)
        
    def add_image(img_path, width_in=5.0):
        if img_path and img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(width_in))
            
    add_heading_1("一、总体概况与核心判断")
    p_desc = (
        f"本次分析共识别 {total} 条有效违章记录，累计记分约 {score_total:.0f} 分。A类或重大违章约 {a_count} 起，"
        f"占比 {high_ratio:.1f}%。数据表明，安全风险不是单点偶发，而是在人员行为、现场执行、协力单位管理和作业许可链条中呈现重复发生特征。"
    )
    add_paragraph(p_desc)
    
    if len(level_counts):
        p_sub = f"从级别结构看，{level_counts.index[0]} 数量最高；从专题归类看，{topic_counts.index[0] if len(topic_counts) else '未知'} 是最突出的高频问题。建议把高频且高后果的专题纳入厂级专项治理。"
        add_paragraph(p_sub)
        
    add_image(trend_path, width_in=5.5)
    add_image(level_path, width_in=4.5)
    
    level_total = float(level_counts.sum()) or 1.0
    level_data = []
    for idx, val in level_counts.items():
        level_data.append([idx, int(val), f"{val / level_total * 100:.1f}%"])
    add_docx_table(doc, ["违章级别", "数量", "占比"], level_data, col_widths=[2.5, 1.5, 1.5])
    
    add_heading_1("二、高频违章与重点区域画像")
    p_topic_desc = (
        f"TOP 专题显示，{topic_counts.index[0] if len(topic_counts) > 0 else '无'}、"
        f"{topic_counts.index[1] if len(topic_counts) > 1 else '其他'} 等问题重复出现，通常对应制度执行弱化、"
        f"班组日常纠偏不足和现场监督穿透力不足。"
    )
    add_paragraph(p_topic_desc)
    
    add_image(topic_path, width_in=5.5)
    add_image(area_path, width_in=5.5)
    
    add_heading_2("高频违章专题 TOP10")
    topic_total = float(topic_counts.sum()) or 1.0
    topic_data = []
    for idx, val in topic_counts.head(10).items():
        topic_data.append([idx, int(val), f"{val / topic_total * 100:.1f}%"])
    add_docx_table(doc, ["违章专题", "数量", "占比"], topic_data, col_widths=[3.0, 1.25, 1.25])
    
    add_heading_2("重点区域/组织 TOP10")
    area_total = float(area_counts.sum()) or 1.0
    area_data = []
    for idx, val in area_counts.head(10).items():
        area_data.append([idx, int(val), f"{val / area_total * 100:.1f}%"])
    add_docx_table(doc, ["区域/组织", "数量", "占比"], area_data, col_widths=[3.0, 1.25, 1.25])
    
    add_heading_1("三、人员性质、协力单位与数据来源分析")
    if not person_counts.empty:
        p_pers = f"人员性质分布显示，{person_counts.index[0]} 违章数量最高，占比 {person_counts.iloc[0] / total * 100:.1f}%。若协力人员占比较高，应重点审视准入培训、作业交底、甲方监护和合同考核闭环。"
        add_paragraph(p_pers)
    if not source_counts.empty:
        p_src = f"数据来源中，{source_counts.index[0]} 占比最高。若行为观察/视频回看占比较高，说明技术监督有效，但也反映现场管理者事中制止不足，需提升现场巡查质量。"
        add_paragraph(p_src)
        
    add_image(person_path, width_in=4.5)
    add_image(source_path, width_in=4.5)
    
    add_heading_1("四、管理体系短板与根因诊断")
    for i, cause in enumerate(root_causes, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{i}. {cause}")
        set_run_font(run, font_name="微软雅黑", size_pt=10.5)
        
    add_heading_1("五、改进建议与行动清单")
    action_data = []
    for action in actions:
        action_data.append([
            action.get('measure', ''),
            action.get('detail', ''),
            action.get('owner', ''),
            action.get('metric', '')
        ])
    add_docx_table(doc, ["措施", "主要做法", "责任主体", "验证指标"], action_data, col_widths=[1.5, 2.5, 1.0, 1.0])
    
    add_heading_1("六、附录：数据字段与口径提示")
    add_paragraph("本报告由 Excel 字段自动统计生成。若源数据含有更细的责任链条、连带记分、隐患编号或整改闭环字段，可进一步扩展为责任追溯、隐患闭环率和整改有效性分析。")
    
    doc.save(docx_path)


def build_pdf(pdf_path: Path, title: str, period: str, total: int, score_total: float, a_count: int, high_ratio: float, 
              level_counts: pd.Series, topic_counts: pd.Series, area_counts: pd.Series, person_counts: pd.Series, source_counts: pd.Series,
              root_causes: list[str], actions: list[dict], 
              trend_path: Path, level_path: Path, topic_path: Path, area_path: Path, person_path: Path, source_path: Path) -> None:
    from matplotlib import font_manager
    
    class SafetyReportPDF(FPDF):
        def header(self):
            if self.page_no() > 1:
                self.set_font("SimHei", size=9)
                self.set_text_color(128, 128, 128)
                self.cell(0, 10, "安全生产违章情况深度分析与改进建议报告", border="B", align="L")
                self.ln(10)
                
        def footer(self):
            self.set_y(-15)
            self.set_font("SimHei", size=9)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"第 {self.page_no()} 页 / {{nb}}", align="C")

    try:
        simhei_path = font_manager.findfont("SimHei")
    except Exception:
        simhei_path = "C:\\Windows\\Fonts\\simhei.ttf"
        
    pdf = SafetyReportPDF()
    pdf.alias_nb_pages()
    
    pdf.add_font("SimHei", style="", fname=simhei_path)
    pdf.add_font("SimHei", style="B", fname=simhei_path)
    
    pdf.add_page()
    
    pdf.set_font("SimHei", style="B", size=18)
    pdf.set_text_color(27, 54, 93)
    pdf.cell(0, 15, title, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    
    pdf.set_font("SimHei", size=9.5)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 5, f"报告日期：{pd.Timestamp.today().strftime('%Y年%m月%d日')}    |    数据周期：{period}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, f"数据口径：有效违章记录 {total} 条，累计记分约 {score_total:.0f} 分，A类/重大违章 {a_count} 起（占比 {high_ratio:.1f}%）", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    
    def add_pdf_heading_1(text):
        pdf.set_font("SimHei", style="B", size=13)
        pdf.set_text_color(27, 54, 93)
        pdf.ln(5)
        pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        
    def add_pdf_heading_2(text):
        pdf.set_font("SimHei", style="B", size=11)
        pdf.set_text_color(27, 54, 93)
        pdf.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
        
    def add_pdf_paragraph(text):
        pdf.set_font("SimHei", size=10)
        pdf.set_text_color(51, 51, 51)
        pdf.multi_cell(0, 6, "    " + text, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
        
    def add_pdf_image(img_path, w=140):
        if img_path and img_path.exists():
            x = (210 - w) / 2
            pdf.image(str(img_path), x=x, w=w)
            pdf.ln(5)

    add_pdf_heading_1("一、总体概况与核心判断")
    p_desc = (
        f"本次分析共识别 {total} 条有效违章记录，累计记分约 {score_total:.0f} 分。A类或重大违章约 {a_count} 起，"
        f"占比 {high_ratio:.1f}%。数据表明，安全风险不是单点偶发，而是在人员行为、现场执行、协力单位管理和作业许可链条中呈现重复发生特征。"
    )
    add_pdf_paragraph(p_desc)
    if len(level_counts):
        p_sub = f"从级别结构看，{level_counts.index[0]} 数量最高；从专题归类看，{topic_counts.index[0] if len(topic_counts) else '未知'} 是最突出的高频问题。建议把高频且高后果的专题纳入厂级专项治理。"
        add_pdf_paragraph(p_sub)
        
    add_pdf_image(trend_path, w=145)
    pdf.ln(2)
    add_pdf_image(level_path, w=110)
    
    level_total = float(level_counts.sum()) or 1.0
    level_headers = ["违章级别", "数量", "占比"]
    with pdf.table(col_widths=(60, 40, 40), text_align="CENTER") as table:
        row = table.row()
        for h in level_headers:
            row.cell(h)
        for idx, val in level_counts.items():
            row = table.row()
            row.cell(str(idx))
            row.cell(str(int(val)))
            row.cell(f"{val / level_total * 100:.1f}%")
            
    pdf.ln(5)
    add_pdf_heading_1("二、高频违章与重点区域画像")
    p_topic_desc = (
        f"TOP 专题显示，{topic_counts.index[0] if len(topic_counts) > 0 else '无'}、"
        f"{topic_counts.index[1] if len(topic_counts) > 1 else '其他'} 等问题重复出现，通常对应制度执行弱化、"
        f"班组日常纠偏不足和现场监督穿透力不足。"
    )
    add_pdf_paragraph(p_topic_desc)
    
    add_pdf_image(topic_path, w=145)
    pdf.add_page()
    add_pdf_image(area_path, w=145)
    
    add_pdf_heading_2("高频违章专题 TOP10")
    topic_total = float(topic_counts.sum()) or 1.0
    with pdf.table(col_widths=(70, 35, 35), text_align="CENTER") as table:
        row = table.row()
        for h in ["违章专题", "数量", "占比"]:
            row.cell(h)
        for idx, val in topic_counts.head(10).items():
            row = table.row()
            row.cell(str(idx))
            row.cell(str(int(val)))
            row.cell(f"{val / topic_total * 100:.1f}%")
            
    pdf.ln(5)
    add_pdf_heading_2("重点区域/组织 TOP10")
    area_total = float(area_counts.sum()) or 1.0
    with pdf.table(col_widths=(70, 35, 35), text_align="CENTER") as table:
        row = table.row()
        for h in ["区域/组织", "数量", "占比"]:
            row.cell(h)
        for idx, val in area_counts.head(10).items():
            row = table.row()
            row.cell(str(idx))
            row.cell(str(int(val)))
            row.cell(f"{val / area_total * 100:.1f}%")
            
    pdf.add_page()
    add_pdf_heading_1("三、人员性质、协力单位与数据来源分析")
    if not person_counts.empty:
        p_pers = f"人员性质分布显示，{person_counts.index[0]} 违章数量最高，占比 {person_counts.iloc[0] / total * 100:.1f}%。若协力人员占比较高，应重点审视准入培训、作业交底、甲方监护和合同考核闭环。"
        add_pdf_paragraph(p_pers)
    if not source_counts.empty:
        p_src = f"数据来源中，{source_counts.index[0]} 占比最高。若行为观察/视频回看占比较高，说明技术监督有效，但也反映现场管理者事中制止不足，需提升现场巡查质量。"
        add_pdf_paragraph(p_src)
        
    add_pdf_image(person_path, w=115)
    add_pdf_image(source_path, w=115)
    
    add_pdf_heading_1("四、管理体系短板与根因诊断")
    pdf.set_font("SimHei", size=10)
    pdf.set_text_color(51, 51, 51)
    for i, cause in enumerate(root_causes, 1):
        pdf.multi_cell(0, 6, f"{i}. {cause}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1.5)
        
    pdf.ln(5)
    add_pdf_heading_1("五、改进建议与行动清单")
    with pdf.table(col_widths=(45, 85, 30, 30), text_align="LEFT") as table:
        row = table.row()
        for h in ["措施", "主要做法", "责任主体", "验证指标"]:
            row.cell(h)
        for action in actions:
            row = table.row()
            row.cell(action.get('measure', ''))
            row.cell(action.get('detail', ''))
            row.cell(action.get('owner', ''))
            row.cell(action.get('metric', ''))
            
    pdf.ln(5)
    add_pdf_heading_1("六、附录：数据字段与口径提示")
    add_pdf_paragraph("本报告由 Excel 字段自动统计生成。若源数据含有更细的责任链条、连带记分、隐患编号或整改闭环字段，可进一步扩展为责任追溯、隐患闭环率和整改有效性分析。")
    
    pdf.output(str(pdf_path))


def build_report(excel_path: Path, output_path: Path, title: str, dynamic_file: Path | None = None) -> Path:
    dynamic_data: dict = {}
    if dynamic_file and dynamic_file.exists():
        try:
            dynamic_data = json.loads(dynamic_file.read_text(encoding="utf-8"))
        except Exception as exc:
            print(f"Warning: Failed to read dynamic file: {exc}")

    df, period = read_export(excel_path)
    if df.empty:
        raise ValueError("清洗后没有有效违章记录。")

    df = df.copy()
    df["专题归类"] = df.apply(classify_topic, axis=1)
    df["简化区域"] = norm_series(df, AREA_COL).map(shorten_area)

    out_dir = output_path.parent / f"{output_path.stem}_charts"
    out_dir.mkdir(parents=True, exist_ok=True)
    trend_path = out_dir / "trend_weekly.png"
    level_path = out_dir / "level_share.png"
    topic_path = out_dir / "topic_top10.png"
    area_path = out_dir / "area_top10.png"
    person_path = out_dir / "person_share.png"
    source_path = out_dir / "source_share.png"

    save_trend(df, trend_path)
    level_counts = norm_series(df, VIOLATION_LEVEL_COL).value_counts()
    topic_counts = norm_series(df, "专题归类").value_counts()
    area_counts = norm_series(df, "简化区域").value_counts()
    person_counts = norm_series(df, PERSON_TYPE_COL).value_counts()
    source_counts = norm_series(df, SOURCE_COL).value_counts()

    save_pie(level_counts, "违章级别占比", level_path)
    save_bar(topic_counts, "高频违章专题 TOP10", topic_path)
    save_bar(area_counts, "重点区域/组织 TOP10", area_path)
    save_pie(person_counts, "违章人员性质占比", person_path)
    save_pie(source_counts, "数据来源占比", source_path)

    total = len(df)
    score_total = df[SCORE_COL].sum() if SCORE_COL in df.columns else 0
    a_count = int(level_counts.get("A类违章", 0) or level_counts.get("A类", 0) or 0)
    high_ratio = a_count / total * 100 if total else 0
    date_min = df[TIME_COL].min().date() if TIME_COL in df.columns and df[TIME_COL].notna().any() else "未识别"
    date_max = df[TIME_COL].max().date() if TIME_COL in df.columns and df[TIME_COL].notna().any() else "未识别"

    root_causes = dynamic_data.get("root_causes") or [
        "风险预控与安全交底流于形式：作业前危险源辨识不充分，交底内容模板化，工票或许可签发未形成真实约束。",
        "核心制度执行衰减：挂牌、联锁、PPE 等低技术门槛要求反复失守，说明班组日常纠偏和管理问责不足。",
        "高风险作业升级管控不足：动火、高处、起重、设备检修等作业未做到许可、确认、监护、复盘闭环。",
        "协力单位穿透式管理不足：甲方对协力队伍的准入、交底、过程监督和绩效联动不足，易形成“以包代管”。",
        "安全培训未转化为行为习惯：员工知道制度但未敬畏风险，说明培训缺少场景化、体验式和实操验证。",
    ]
    actions = dynamic_data.get("actions") or [
        {
            "measure": "立即开展高频问题专项整治",
            "detail": "围绕挂牌/能量隔离、PPE、动火、高处、起重等专题开展 1-3 个月专项行动，执行零容忍停工整改。",
            "owner": "厂部/分厂/作业区/协力单位",
            "metric": "违章数下降、A类清零、重复违章减少、闭环率100%",
        },
        {
            "measure": "重构高风险作业许可",
            "detail": "许可签发人、作业负责人、监护人必须现场逐项确认并拍照留痕；关键风险点使用清单化确认。",
            "owner": "厂部/分厂/作业区/协力单位",
            "metric": "许可执行率提升、违规作业显著下降",
        },
        {
            "measure": "强化管理人员履职量化",
            "detail": "建立班组长、作业长、分厂管理者安全履职清单，将制止违章、现场巡查、交底审核与绩效挂钩。",
            "owner": "厂部/分厂/作业区",
            "metric": "现场纠偏频次提升、重复问题下降",
        },
        {
            "measure": "实施协力单位等同管理",
            "detail": "协力人员参加同等培训和班前会；月度发布协力单位安全积分，并与合同结算、清退机制联动。",
            "owner": "厂部/采购/协力单位",
            "metric": "协力违章占比下降、准入合规率提升",
        },
        {
            "measure": "推进工程技术防呆",
            "detail": "对钢包、剪切、步进梁、小车运行区等高风险点推广硬隔离、联锁、权限钥匙和感应停机。",
            "owner": "设备/工艺/安全部门",
            "metric": "同类机械伤害风险下降",
        },
        {
            "measure": "建立数据驱动闭环",
            "detail": "每周更新违章看板，每月组织跨部门复盘，针对异常升高专题形成整改责任书并跟踪验证。",
            "owner": "安全管理部门",
            "metric": "整改按期完成率提升、异常专题响应时效缩短",
        },
    ]

    lines: list[str] = [
        f"# {title}",
        "",
        f"💡 **提示**：本分析已同步生成专业格式的报告文档，您可直接下载：",
        f"- 📝 **[点击下载 Word 报告](./{output_path.name.replace('.md', '.docx')})**",
        f"- 📄 **[点击下载 PDF 报告](./{output_path.name.replace('.md', '.pdf')})**",
        "",
        f"- 报告日期：{pd.Timestamp.today().strftime('%Y年%m月%d日')}",
        f"- 数据周期：{period or f'{date_min} 至 {date_max}'}",
        f"- 数据口径：剔除“已删除”记录，共 {total} 条有效违章明细；如源文件为“不含连带”导出，则连带责任仅作定性参考。",
        "",
        "## 一、总体概况与核心判断",
        "",
        f"本次分析共识别 {total} 条有效违章记录，累计记分约 {score_total:.0f} 分。A类或重大违章约 {a_count} 起，占比 {high_ratio:.1f}%。数据表明，安全风险不是单点偶发，而是在人员行为、现场执行、协力单位管理和作业许可链条中呈现重复发生特征。",
        f"从级别结构看，{level_counts.index[0] if len(level_counts) else '未知'} 数量最高；从专题归类看，{topic_counts.index[0] if len(topic_counts) else '未知'} 是最突出的高频问题。建议把高频且高后果的专题纳入厂级专项治理。",
        "",
        f"![违章发生趋势]({relative_asset_path(output_path, trend_path)})",
        "",
        f"![违章级别占比]({relative_asset_path(output_path, level_path)})",
        "",
        *format_count_table(level_counts),
        "",
        "## 二、高频违章与重点区域画像",
        "",
        f"TOP 专题显示，{topic_counts.index[0] if len(topic_counts) > 0 else '无'}、{topic_counts.index[1] if len(topic_counts) > 1 else '其他'} 等问题重复出现，通常对应制度执行弱化、班组日常纠偏不足和现场监督穿透力不足。",
        "",
        f"![高频违章专题 TOP10]({relative_asset_path(output_path, topic_path)})",
        "",
        f"![重点区域/组织 TOP10]({relative_asset_path(output_path, area_path)})",
        "",
        "### 高频违章专题 TOP10",
        "",
        *format_count_table(topic_counts),
        "",
        "### 重点区域/组织 TOP10",
        "",
        *format_count_table(area_counts),
        "",
        "## 三、人员性质、协力单位与数据来源分析",
        "",
    ]

    if not person_counts.empty:
        lines.extend([
            f"人员性质分布显示，{person_counts.index[0]} 违章数量最高，占比 {person_counts.iloc[0] / total * 100:.1f}%。若协力人员占比较高，应重点审视准入培训、作业交底、甲方监护和合同考核闭环。",
            "",
        ])
    if not source_counts.empty:
        lines.extend([
            f"数据来源中，{source_counts.index[0]} 占比最高。若行为观察/视频回看占比较高，说明技术监督有效，但也反映现场管理者事中制止不足，需提升现场巡查质量。",
            "",
        ])

    lines.extend([
        f"![违章人员性质占比]({relative_asset_path(output_path, person_path)})",
        "",
        f"![数据来源占比]({relative_asset_path(output_path, source_path)})",
        "",
        "## 四、管理体系短板与根因诊断",
        "",
    ])
    lines.extend([f"1. {item}" for item in root_causes])
    lines.extend([
        "",
        "## 五、改进建议与行动清单",
        "",
        "| 措施 | 主要做法 | 责任主体 | 验证指标 |",
        "| --- | --- | --- | --- |",
    ])
    for action in actions:
        lines.append(
            f"| {action.get('measure', '')} | {action.get('detail', '')} | {action.get('owner', '')} | {action.get('metric', '')} |"
        )

    lines.extend([
        "",
        "## 六、附录：数据字段与口径提示",
        "",
        "本报告由 Excel 字段自动统计生成。若源数据含有更细的责任链条、连带记分、隐患编号或整改闭环字段，可进一步扩展为责任追溯、隐患闭环率和整改有效性分析。",
        "",
    ])

    output_path.write_text("\n".join(lines), encoding="utf-8")

    docx_path = output_path.with_suffix(".docx")
    pdf_path = output_path.with_suffix(".pdf")

    try:
        build_docx(
            docx_path=docx_path,
            title=title,
            period=period or f"{date_min} 至 {date_max}",
            total=total,
            score_total=score_total,
            a_count=a_count,
            high_ratio=high_ratio,
            level_counts=level_counts,
            topic_counts=topic_counts,
            area_counts=area_counts,
            person_counts=person_counts,
            source_counts=source_counts,
            root_causes=root_causes,
            actions=actions,
            trend_path=trend_path,
            level_path=level_path,
            topic_path=topic_path,
            area_path=area_path,
            person_path=person_path,
            source_path=source_path
        )
        print(f"Word report successfully generated at: {docx_path}")
    except Exception as e:
        print(f"Error generating Word report: {e}")
        import traceback
        traceback.print_exc()

    try:
        build_pdf(
            pdf_path=pdf_path,
            title=title,
            period=period or f"{date_min} 至 {date_max}",
            total=total,
            score_total=score_total,
            a_count=a_count,
            high_ratio=high_ratio,
            level_counts=level_counts,
            topic_counts=topic_counts,
            area_counts=area_counts,
            person_counts=person_counts,
            source_counts=source_counts,
            root_causes=root_causes,
            actions=actions,
            trend_path=trend_path,
            level_path=level_path,
            topic_path=topic_path,
            area_path=area_path,
            person_path=person_path,
            source_path=source_path
        )
        print(f"PDF report successfully generated at: {pdf_path}")
    except Exception as e:
        print(f"Error generating PDF report: {e}")
        import traceback
        traceback.print_exc()

    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate safety violation analysis report from Excel.")
    parser.add_argument("excel", type=Path, help="违章导出 Excel 文件路径")
    parser.add_argument("--output", "-o", type=Path, default=Path("safety_violation_report.md"), help="输出 Markdown 路径")
    parser.add_argument("--title", default="安全生产违章情况深度分析与改进建议报告")
    parser.add_argument("--export-stats", action="store_true", help="只输出统计数据的 JSON")
    parser.add_argument("--dynamic-file", type=Path, help="包含动态分析内容的 JSON 文件路径")
    args = parser.parse_args()

    if args.export_stats:
        export_stats(args.excel)
    else:
        build_report(args.excel, args.output, args.title, args.dynamic_file)
        print(args.output)


if __name__ == "__main__":
    main()

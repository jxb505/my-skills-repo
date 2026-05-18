#!/usr/bin/env python3
"""Generate a Chinese safety-production violation analysis DOCX with charts from an Excel export."""
from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

EXPECTED = ["状态", "违章区域", "检查机构", "检查人员类别", "违章时间", "违章人员性质", "总包单位", "分包单位", "违章情况", "违章类别", "违章级别", "分值", "数据来源", "条款内容", "连带记分信息"]


def setup_chinese_font() -> None:
    candidates = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Noto Sans CJK JP", "WenQuanYi Micro Hei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["font.sans-serif"] = candidates
    plt.rcParams["axes.unicode_minus"] = False


def read_export(path: Path) -> tuple[pd.DataFrame, str]:
    raw = pd.read_excel(path, sheet_name=0, header=None)
    period = ""
    for value in raw.stack().dropna().astype(str).head(50):
        if "时间范围" in value:
            period = value.replace("时间范围：", "").replace("时间范围:", "")
            break
    header_idx = None
    for i in range(min(20, len(raw))):
        row = [str(x).strip() for x in raw.iloc[i].fillna("").tolist()]
        if "状态" in row and "违章情况" in row:
            header_idx = i
            break
    if header_idx is None:
        raise ValueError("无法识别表头行；请确认 Excel 包含“状态”“违章情况”等字段。")
    df = pd.read_excel(path, sheet_name=0, header=header_idx)
    df = df.loc[:, ~df.columns.astype(str).str.startswith("Unnamed")]
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all")
    if "状态" in df.columns:
        df = df[df["状态"].astype(str).str.strip().ne("已删除")]
    if "违章时间" in df.columns:
        df["违章时间"] = pd.to_datetime(df["违章时间"], errors="coerce")
    if "分值" in df.columns:
        df["分值"] = pd.to_numeric(df["分值"], errors="coerce").fillna(0)
    return df, period


def norm_series(df: pd.DataFrame, col: str) -> pd.Series:
    if col not in df.columns:
        return pd.Series(dtype=object)
    return df[col].fillna("未填写").astype(str).str.strip().replace({"": "未填写", "nan": "未填写"})


def shorten_area(text: str) -> str:
    parts = [p for p in str(text).split("/") if p]
    if len(parts) >= 3:
        return "/".join(parts[-3:])
    return str(text)


def classify_topic(row: pd.Series) -> str:
    text = " ".join(str(row.get(c, "")) for c in ["违章情况", "条款内容", "违章类别"])
    rules = [
        ("三方挂牌/能量隔离", r"挂牌|摘牌|锁定|联锁|能量|确认"),
        ("个体防护/PPE", r"防护|护目镜|面罩|耳塞|口罩|手套|安全帽|安全带|围裙|劳防"),
        ("高处/临边", r"高处|登高|临边|孔洞|坠落|安全绳|护栏"),
        ("动火/消防", r"动火|火星|灭火器|接火|可燃|氧气|乙炔|消防"),
        ("起重吊装", r"起重|吊装|吊物|吊索|司索|指吊|行车|歪拉|斜吊"),
        ("安全交底/工票", r"交底|工票|签字|代签|审批|许可"),
        ("设备设施/机械伤害", r"设备|剪|小车|步进梁|打捆|机械|旋转|挤压"),
        ("行为规范/禁令", r"吸烟|酒|手机|禁令|无证|准入|睡岗"),
    ]
    for label, pattern in rules:
        if re.search(pattern, text):
            return label
    return str(row.get("违章类别", "其他")) or "其他"


def save_bar(series: pd.Series, title: str, path: Path, xlabel: str = "数量") -> None:
    setup_chinese_font()
    data = series.dropna().head(10).sort_values()
    fig, ax = plt.subplots(figsize=(8, max(4, 0.45 * len(data))))
    ax.barh(data.index.astype(str), data.values)
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    for i, v in enumerate(data.values):
        ax.text(v, i, f" {int(v)}", va="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def save_pie(series: pd.Series, title: str, path: Path) -> None:
    setup_chinese_font()
    data = series.dropna()
    fig, ax = plt.subplots(figsize=(6.2, 4.6))
    ax.pie(data.values, labels=data.index.astype(str), autopct="%1.1f%%", startangle=90)
    ax.set_title(title)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def save_trend(df: pd.DataFrame, path: Path) -> pd.Series:
    setup_chinese_font()
    if "违章时间" not in df.columns or df["违章时间"].isna().all():
        trend = pd.Series(dtype=int)
    else:
        trend = df.dropna(subset=["违章时间"]).set_index("违章时间").resample("W").size()
    fig, ax = plt.subplots(figsize=(8, 4.2))
    if len(trend):
        ax.plot(trend.index, trend.values, marker="o")
        ax.set_xticklabels([d.strftime("%m-%d") for d in trend.index], rotation=35, ha="right")
    ax.set_title("违章发生趋势（按周）")
    ax.set_ylabel("违章数量")
    ax.grid(True, alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return trend


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25


def add_table_from_series(doc: Document, title: str, series: pd.Series, value_name: str = "数量") -> None:
    doc.add_paragraph(title).runs[0].bold = True
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "项目", value_name, "占比"
    total = float(series.sum()) or 1.0
    for idx, val in series.head(10).items():
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(int(val)) if float(val).is_integer() else f"{val:.1f}"
        row[2].text = f"{val / total * 100:.1f}%"


def export_stats(excel_path: Path) -> None:
    df, period = read_export(excel_path)
    if df.empty:
        print(json.dumps({"error": "清洗后没有有效违章记录"}, ensure_ascii=False))
        return

    df["专题归类"] = df.apply(classify_topic, axis=1)
    df["简化区域"] = norm_series(df, "违章区域").map(shorten_area)

    level_counts = norm_series(df, "违章级别").value_counts()
    topic_counts = norm_series(df, "专题归类").value_counts()
    area_counts = norm_series(df, "简化区域").value_counts()

    total = len(df)
    score_total = df["分值"].sum() if "分值" in df.columns else 0
    a_count = int(level_counts.get("A类违章", 0) or level_counts.get("A类", 0) or 0)

    stats = {
        "period": period,
        "total_records": total,
        "total_score": float(score_total),
        "a_level_count": a_count,
        "top_topics": {str(k): int(v) for k, v in topic_counts.head(5).items()},
        "top_areas": {str(k): int(v) for k, v in area_counts.head(5).items()}
    }
    print(json.dumps(stats, ensure_ascii=False, indent=2))


def build_report(excel_path: Path, output_path: Path, title: str, dynamic_file: Path | None = None) -> Path:
    dynamic_data = {}
    if dynamic_file and dynamic_file.exists():
        try:
            with open(dynamic_file, "r", encoding="utf-8") as f:
                dynamic_data = json.load(f)
        except Exception as e:
            print(f"Warning: Failed to read dynamic file: {e}")

    df, period = read_export(excel_path)
    if df.empty:
        raise ValueError("清洗后没有有效违章记录。")
    df = df.copy()
    df["专题归类"] = df.apply(classify_topic, axis=1)
    df["简化区域"] = norm_series(df, "违章区域").map(shorten_area)

    out_dir = output_path.parent / (output_path.stem + "_charts")
    out_dir.mkdir(parents=True, exist_ok=True)
    trend_path = out_dir / "trend_weekly.png"
    level_path = out_dir / "level_share.png"
    topic_path = out_dir / "topic_top10.png"
    area_path = out_dir / "area_top10.png"
    person_path = out_dir / "person_share.png"
    source_path = out_dir / "source_share.png"

    trend = save_trend(df, trend_path)
    level_counts = norm_series(df, "违章级别").value_counts()
    topic_counts = norm_series(df, "专题归类").value_counts()
    area_counts = norm_series(df, "简化区域").value_counts()
    person_counts = norm_series(df, "违章人员性质").value_counts()
    source_counts = norm_series(df, "数据来源").value_counts()
    save_pie(level_counts, "违章级别占比", level_path)
    save_bar(topic_counts, "高频违章专题 TOP10", topic_path)
    save_bar(area_counts, "重点区域/组织 TOP10", area_path)
    save_pie(person_counts, "违章人员性质占比", person_path)
    save_pie(source_counts, "数据来源占比", source_path)

    total = len(df)
    score_total = df["分值"].sum() if "分值" in df.columns else 0
    a_count = int(level_counts.get("A类违章", 0) or level_counts.get("A类", 0) or 0)
    high_ratio = a_count / total * 100 if total else 0
    date_min = df["违章时间"].min().date() if "违章时间" in df.columns and df["违章时间"].notna().any() else "未识别"
    date_max = df["违章时间"].max().date() if "违章时间" in df.columns and df["违章时间"].notna().any() else "未识别"

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7); sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(18)
    doc.add_paragraph(f"报告日期：{pd.Timestamp.today().strftime('%Y年%m月%d日')}")
    doc.add_paragraph(f"数据周期：{period or f'{date_min} 至 {date_max}'}")
    doc.add_paragraph(f"数据口径：剔除“已删除”记录，共 {total} 条有效违章明细；如源文件为“不含连带”导出，则连带责任仅作定性参考。")

    add_heading(doc, "一、总体概况与核心判断", 1)
    add_para(doc, f"本次分析共识别 {total} 条有效违章记录，累计记分约 {score_total:.0f} 分。A类或重大违章约 {a_count} 起，占比 {high_ratio:.1f}%。数据表明，安全风险不是单点偶发，而是在人员行为、现场执行、协力单位管理和作业许可链条中呈现重复发生特征。")
    add_para(doc, f"从级别结构看，{level_counts.index[0] if len(level_counts) else '未知'} 数量最高；从专题归类看，{topic_counts.index[0] if len(topic_counts) else '未知'} 是最突出的高频问题。建议把高频且高后果的专题纳入厂级专项治理。")
    doc.add_picture(str(trend_path), width=Inches(6.6))
    doc.add_picture(str(level_path), width=Inches(5.8))
    add_table_from_series(doc, "违章级别分布", level_counts)

    add_heading(doc, "二、高频违章与重点区域画像", 1)
    add_para(doc, f"TOP 专题显示，{topic_counts.index[0] if len(topic_counts)>0 else '无'}、{topic_counts.index[1] if len(topic_counts)>1 else '其他'} 等问题重复出现，通常对应制度执行弱化、班组日常纠偏不足和现场监督穿透力不足。")
    doc.add_picture(str(topic_path), width=Inches(6.6))
    doc.add_picture(str(area_path), width=Inches(6.6))
    add_table_from_series(doc, "高频违章专题 TOP10", topic_counts)
    add_table_from_series(doc, "重点区域/组织 TOP10", area_counts)

    add_heading(doc, "三、人员性质、协力单位与数据来源分析", 1)
    if not person_counts.empty:
        add_para(doc, f"人员性质分布显示，{person_counts.index[0]} 违章数量最高，占比 {person_counts.iloc[0]/total*100:.1f}%。若协力人员占比较高，应重点审视准入培训、作业交底、甲方监护和合同考核闭环。")
    if not source_counts.empty:
        add_para(doc, f"数据来源中，{source_counts.index[0]} 占比最高。若行为观察/视频回看占比较高，说明技术监督有效，但也反映现场管理者事中制止不足，需提升现场巡查质量。")
    doc.add_picture(str(person_path), width=Inches(5.8))
    doc.add_picture(str(source_path), width=Inches(5.8))

    add_heading(doc, "四、管理体系短板与根因诊断", 1)
    root_causes = dynamic_data.get("root_causes")
    if not root_causes:
        root_causes = [
            "风险预控与安全交底流于形式：作业前危险源辨识不充分，交底内容模板化，工票或许可签发未形成真实约束。",
            "核心制度执行衰减：挂牌、联锁、PPE 等低技术门槛要求反复失守，说明班组日常纠偏和管理问责不足。",
            "高风险作业升级管控不足：动火、高处、起重、设备检修等作业未做到许可、确认、监护、复盘闭环。",
            "协力单位穿透式管理不足：甲方对协力队伍的准入、交底、过程监护和绩效联动不足，易形成“以包代管”。",
            "安全培训未转化为行为习惯：员工知道制度但未敬畏风险，说明培训缺少场景化、体验式和实操验证。",
        ]
    for item in root_causes:
        add_para(doc, item)

    add_heading(doc, "五、改进建议与行动清单", 1)
    actions = dynamic_data.get("actions")
    if not actions:
        actions = [
            {"measure": "立即开展高频顽疾专项整治", "detail": "围绕挂牌/能量隔离、PPE、动火、高处、起重等专题开展 1-3 个月专项行动，执行零容忍停工整改。", "owner": "厂部/分厂/作业区/协力单位", "metric": "违章数下降、A类清零、重复违章减少、闭环率100%"},
            {"measure": "重构高风险作业许可", "detail": "许可签发人、作业负责人、监护人必须现场逐项确认并拍照留痕；关键风险点使用清单化确认。", "owner": "厂部/分厂/作业区/协力单位", "metric": "违章数下降、A类清零、重复违章减少、闭环率100%"},
            {"measure": "强化管理人员履职量化", "detail": "建立班组长、作业长、分厂管理者安全履职清单，将制止违章、现场巡查、交底审核与绩效挂钩。", "owner": "厂部/分厂/作业区/协力单位", "metric": "违章数下降、A类清零、重复违章减少、闭环率100%"},
            {"measure": "实施协力单位等同管理", "detail": "协力人员参加同等培训和班前会；月度发布协力单位安全积分，和合同结算、清退机制联动。", "owner": "厂部/分厂/作业区/协力单位", "metric": "违章数下降、A类清零、重复违章减少、闭环率100%"},
            {"measure": "推进工程技术防御", "detail": "对锌锅、剪区、步进梁、小车运行区等高风险点推广硬隔离、联锁、权限钥匙和感应停机。", "owner": "厂部/分厂/作业区/协力单位", "metric": "违章数下降、A类清零、重复违章减少、闭环率100%"},
            {"measure": "建立数据驱动闭环", "detail": "每周更新违章看板，每月组织跨部门复盘，针对异常升高专题形成整改责任书并跟踪验证。", "owner": "厂部/分厂/作业区/协力单位", "metric": "违章数下降、A类清零、重复违章减少、闭环率100%"},
        ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["措施", "主要做法", "责任主体", "验证指标"]):
        cell.text = text
    for action in actions:
        row = table.add_row().cells
        row[0].text = action.get("measure", "")
        row[1].text = action.get("detail", "")
        row[2].text = action.get("owner", "")
        row[3].text = action.get("metric", "")

    add_heading(doc, "六、附录：数据字段与口径提示", 1)
    add_para(doc, "本报告由 Excel 字段自动统计生成。若源数据含有更细的责任链条、连带记分、隐患编号或整改闭环字段，可进一步扩展为责任追溯、隐患闭环率和整改有效性分析。")
    doc.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate safety violation analysis report from Excel.")
    parser.add_argument("excel", type=Path, help="违章导出 Excel 文件路径")
    parser.add_argument("--output", "-o", type=Path, default=Path("safety_violation_report.docx"), help="输出 DOCX 路径")
    parser.add_argument("--title", default="安全生产违章情况深度分析与改进建议报告")
    parser.add_argument("--export-stats", action="store_true", help="只输出统计数据的 JSON")
    parser.add_argument("--dynamic-file", type=Path, help="包含动态思考内容的 JSON 文件路径")
    args = parser.parse_args()

    if args.export_stats:
        export_stats(args.excel)
    else:
        build_report(args.excel, args.output, args.title, args.dynamic_file)
        print(args.output)


if __name__ == "__main__":
    main()
